import json
import time
import io
import docx
import streamlit as st
import streamlit.components.v1 as components
from PIL import Image
from google import genai
from google.genai import types


# =============================================================================
# 1. PAGE CONFIGURATION
# =============================================================================

st.set_page_config(
    page_title="Pitch to Project | Smart Scope Engine",
    page_icon="⚡",
    layout="wide",
)


# =============================================================================
# 2. GLOBAL NEON CSS
# =============================================================================

css_code = """
<style>

/* -------------------------------------------------------------------------
   GLOBAL APPLICATION BACKGROUND
   ------------------------------------------------------------------------- */

.stApp {
    background:
        linear-gradient(
            125deg,
            #0f172a 0%,
            #1e1b4b 35%,
            #311042 70%,
            #0284c7 100%
        ) !important;

    background-attachment: fixed;
}


/* -------------------------------------------------------------------------
   LOGIN FORM
   ------------------------------------------------------------------------- */

div[data-testid="stForm"] {
    background: rgba(15, 23, 42, 0.95) !important;
    border: 2.5px solid #a855f7 !important;
    border-radius: 18px !important;
    padding: 36px !important;
    box-shadow: 0 0 40px rgba(168, 85, 247, 0.5) !important;
}


/* -------------------------------------------------------------------------
   LOGIN LABELS
   ------------------------------------------------------------------------- */

div[data-testid="stForm"] label,
div[data-testid="stForm"] label p,
div[data-testid="stTextInput"] label p {
    color: #38bdf8 !important;
    font-weight: 900 !important;
    font-size: 1.4rem !important;
    letter-spacing: 0.5px !important;
    margin-bottom: 8px !important;
    text-align: left !important;
    width: 100% !important;
    display: block !important;
}


/* -------------------------------------------------------------------------
   INPUT FIELDS
   ------------------------------------------------------------------------- */

div[data-testid="stForm"] div[data-testid="stTextInput"] input,
div[data-testid="stTextInput"] input,
input[type="text"],
input[type="password"] {

    background-color: #0f172a !important;
    color: #ffffff !important;
    -webkit-text-fill-color: #ffffff !important;

    border: 2.5px solid #38bdf8 !important;
    border-radius: 12px !important;

    font-weight: 800 !important;
    font-size: 1.35rem !important;

    padding: 16px 20px !important;
    text-align: left !important;

    box-shadow:
        0 0 14px rgba(56, 189, 248, 0.3) !important;
}


/* -------------------------------------------------------------------------
   BROWSER AUTOFILL
   ------------------------------------------------------------------------- */

input:-webkit-autofill,
input:-webkit-autofill:hover,
input:-webkit-autofill:focus,
input:-webkit-autofill:active {

    -webkit-box-shadow:
        0 0 0 1000px #0f172a inset !important;

    -webkit-text-fill-color: #ffffff !important;

    caret-color: #ffffff !important;

    border: 2.5px solid #38bdf8 !important;
}


/* -------------------------------------------------------------------------
   PLACEHOLDER
   ------------------------------------------------------------------------- */

div[data-testid="stTextInput"] input::placeholder,
input::placeholder {

    color: #94a3b8 !important;
    -webkit-text-fill-color: #94a3b8 !important;

    font-weight: 700 !important;
    font-size: 1.25rem !important;

    opacity: 1 !important;
    text-align: left !important;
}


/* -------------------------------------------------------------------------
   PASSWORD EYE ICON
   ------------------------------------------------------------------------- */

div[data-testid="stTextInput"] button svg,
div[data-testid="stForm"] svg {

    fill: #38bdf8 !important;
    stroke: #38bdf8 !important;

    width: 26px !important;
    height: 26px !important;
}


/* -------------------------------------------------------------------------
   LOGIN BUTTON
   ------------------------------------------------------------------------- */

div[data-testid="stForm"] button[type="submit"],
div[data-testid="stForm"] button[data-testid="stFormSubmitButton"],
div[data-testid="stForm"] button {

    background:
        linear-gradient(
            90deg,
            #ec4899 0%,
            #8b5cf6 50%,
            #06b6d4 100%
        ) !important;

    border: none !important;
    border-radius: 14px !important;

    padding: 18px 32px !important;

    box-shadow:
        0 0 30px rgba(236, 72, 153, 0.7) !important;

    transition: all 0.3s ease !important;

    margin-top: 22px !important;

    width: 100% !important;
}


div[data-testid="stForm"] button[type="submit"] *,
div[data-testid="stForm"] button[type="submit"] p,
div[data-testid="stForm"] button[type="submit"] span,
div[data-testid="stForm"] button[data-testid="stFormSubmitButton"] *,
div[data-testid="stForm"] button[data-testid="stFormSubmitButton"] p,
div[data-testid="stForm"] button[data-testid="stFormSubmitButton"] span {

    color: #ffffff !important;
    -webkit-text-fill-color: #ffffff !important;

    font-weight: 900 !important;
    font-size: 1.5rem !important;

    letter-spacing: 1.2px !important;
}


/* -------------------------------------------------------------------------
   SIDEBAR
   ------------------------------------------------------------------------- */

section[data-testid="stSidebar"] {

    background: rgba(15, 23, 42, 0.95) !important;

    border-right:
        1.5px solid #a855f7 !important;
}


section[data-testid="stSidebar"] h3 {

    color: #38bdf8 !important;

    font-size: 1.3rem !important;

    font-weight: 800 !important;
}


section[data-testid="stSidebar"] p,
section[data-testid="stSidebar"] span,
section[data-testid="stSidebar"] div {

    color: #f8fafc !important;

    font-size: 1.05rem !important;

    font-weight: 700 !important;

    background: transparent !important;
}


section[data-testid="stSidebar"] div.stButton > button {

    background:
        linear-gradient(
            90deg,
            #ec4899 0%,
            #f43f5e 100%
        ) !important;

    color: #ffffff !important;

    font-weight: 800 !important;

    font-size: 1.1rem !important;

    border-radius: 10px !important;

    box-shadow:
        0 0 15px rgba(244, 63, 94, 0.5) !important;

    border: none !important;

    margin-top: 10px !important;
}


/* -------------------------------------------------------------------------
   HEADINGS
   ------------------------------------------------------------------------- */

h2,
h3 {

    color: #f8fafc !important;

    font-weight: 800 !important;

    text-shadow:
        0 0 10px rgba(168, 85, 247, 0.3) !important;
}


div[data-testid="stMarkdownContainer"] p,
label[data-testid="stWidgetLabel"] p,
div[data-testid="stToggle"] span {

    color: #f8fafc !important;

    font-weight: 700 !important;
}


/* -------------------------------------------------------------------------
   FILE UPLOADER
   ------------------------------------------------------------------------- */

div[data-testid="stFileUploader"] {

    background:
        rgba(15, 23, 42, 0.95) !important;

    border:
        2px solid #a855f7 !important;

    border-radius: 14px !important;

    padding: 12px !important;
}


div[data-testid="stFileUploader"] section,
div[data-testid="stFileUploaderDropzone"],
div[data-testid="stFileUploader"]
[data-testid="stFileUploaderDropzone"] {

    background: #1e1b4b !important;

    border:
        2px dashed #38bdf8 !important;

    border-radius: 10px !important;
}


div[data-testid="stFileUploaderDropzone"] *,
div[data-testid="stFileUploaderDropzone"] span,
div[data-testid="stFileUploaderDropzone"] small,
div[data-testid="stFileUploaderDropzone"] p {

    color: #ffffff !important;

    font-weight: 700 !important;
}


/* -------------------------------------------------------------------------
   TOAST
   ------------------------------------------------------------------------- */

div[data-testid="stToast"],
div[data-testid="stToast"] > div {

    background-color: #1e1b4b !important;
    background: #1e1b4b !important;

    border:
        2px solid #38bdf8 !important;

    border-radius: 12px !important;

    box-shadow:
        0 0 20px rgba(56, 189, 248, 0.5) !important;
}


div[data-testid="stToast"] * {

    color: #ffffff !important;

    font-weight: 800 !important;
}


/* -------------------------------------------------------------------------
   DOWNLOAD BUTTON
   ------------------------------------------------------------------------- */

div[data-testid="stDownloadButton"] > button {

    background:
        linear-gradient(
            135deg,
            #10b981 0%,
            #059669 100%
        ) !important;

    border-radius: 12px !important;

    border: none !important;

    box-shadow:
        0 0 18px rgba(16, 185, 129, 0.5) !important;

    width: 100%;
}


div[data-testid="stDownloadButton"] > button * {

    color: #ffffff !important;

    font-weight: 900 !important;
}


/* -------------------------------------------------------------------------
   TEXT AREA
   ------------------------------------------------------------------------- */

div[data-testid="stTextArea"] textarea {

    background:
        rgba(15, 23, 42, 0.85) !important;

    border:
        2px solid #a855f7 !important;

    border-radius: 14px !important;

    color: #ffffff !important;
}


/* -------------------------------------------------------------------------
   NORMAL BUTTONS
   ------------------------------------------------------------------------- */

div.stButton > button {

    background:
        linear-gradient(
            90deg,
            #ec4899 0%,
            #8b5cf6 50%,
            #3b82f6 100%
        ) !important;

    color: #ffffff !important;

    font-weight: 800 !important;

    font-size: 1.05rem !important;

    border-radius: 12px !important;

    border: none !important;

    padding: 14px 28px !important;

    box-shadow:
        0 0 20px rgba(139, 92, 246, 0.5) !important;

    width: 100%;
}


/* -------------------------------------------------------------------------
   SELECTED TABS
   ------------------------------------------------------------------------- */

button[aria-selected="true"] {

    background:
        linear-gradient(
            135deg,
            #8b5cf6 0%,
            #ec4899 100%
        ) !important;

    color: #ffffff !important;
}


/* -------------------------------------------------------------------------
   CARDS
   ------------------------------------------------------------------------- */

div[data-testid="stVerticalBlockBorderWrapper"] > div {

    background:
        rgba(15, 23, 42, 0.8) !important;

    backdrop-filter: blur(10px) !important;

    border-left:
        6px solid #06b6d4 !important;

    border-radius: 14px !important;

    padding: 20px !important;
}


/* -------------------------------------------------------------------------
   CODE
   ------------------------------------------------------------------------- */

code {

    background-color:
        rgba(30, 27, 75, 0.95) !important;

    color: #38bdf8 !important;

    border:
        1px solid #a855f7 !important;

    border-radius: 6px !important;

    padding: 3px 8px !important;
}

</style>
"""

st.markdown(css_code, unsafe_allow_html=True)


# =============================================================================
# 3. AUTHENTICATION
# =============================================================================

def check_password():
    """
    Returns True if the user enters valid credentials.
    """

    if st.session_state.get("authenticated", False):
        return True

    st.markdown("<br><br>", unsafe_allow_html=True)

    col1, col2, col3 = st.columns([1, 2.4, 1])

    with col2:

        with st.form("login_form"):

            components.html(
                """
                <div style="
                    background:
                        linear-gradient(
                            135deg,
                            #ec4899 0%,
                            #8b5cf6 50%,
                            #06b6d4 100%
                        );

                    border-radius: 14px;

                    padding: 24px 10px;

                    box-shadow:
                        0 0 25px rgba(236, 72, 153, 0.6);

                    text-align: center;

                    font-family:
                        system-ui,
                        -apple-system,
                        sans-serif;
                ">

                    <div style="
                        color: #ffffff;

                        font-size: 2.8rem;

                        font-weight: 900;

                        margin-bottom: 6px;

                        text-shadow:
                            0 3px 12px rgba(0, 0, 0, 0.8);

                        letter-spacing: -0.5px;
                    ">
                        🔒 Pitch to Project
                    </div>

                    <div style="
                        color: #ffffff;

                        font-size: 1.35rem;

                        font-weight: 800;

                        letter-spacing: 1px;

                        text-shadow:
                            0 2px 8px rgba(0, 0, 0, 0.8);
                    ">
                        ⚡ Scope Intelligence Engine Access
                    </div>

                </div>
                """,
                height=140,
                scrolling=False
            )

            username = st.text_input(
                "Username",
                placeholder="Enter username"
            )

            password = st.text_input(
                "Password",
                type="password",
                placeholder="Enter password"
            )

            submit = st.form_submit_button(
                "🔑 LOGIN TO ENGINE"
            )

            if submit:

                valid_user = st.secrets.get(
                    "APP_USER",
                    "adin"
                )

                valid_password = st.secrets.get(
                    "APP_PASSWORD",
                    "project2026"
                )

                if (
                    username == valid_user
                    and password == valid_password
                ):

                    st.session_state["authenticated"] = True

                    st.toast(
                        "⚡ Login Successful!",
                        icon="✅"
                    )

                    st.rerun()

                else:

                    st.error(
                        "❌ Invalid Username or Password"
                    )

    return False


# Stop application until authenticated
if not check_password():
    st.stop()


# =============================================================================
# 4. SIDEBAR SESSION
# =============================================================================

with st.sidebar:

    st.markdown("### 👤 User Session")

    st.write("Logged in as **Admin**")

    if st.button(
        "🚪 Logout",
        use_container_width=True
    ):

        st.session_state["authenticated"] = False

        st.rerun()


# =============================================================================
# 5. GEMINI CLIENT
# =============================================================================

@st.cache_resource
def get_gemini_client():

    try:

        api_key = st.secrets["GEMINI_API_KEY"]

        return genai.Client(
            api_key=api_key
        )

    except Exception:

        return None


client = get_gemini_client()


# =============================================================================
# 6. PROGRESS BAR (FIXED CSS CONTAINER COLLAPSE ISSUE)
# =============================================================================

def render_stylish_progress(percentage, status_text):
    return f"""
    <div style="margin: 10px 0 18px 0; width: 100%;">
        <div style="
            display: flex;
            justify-content: space-between;
            align-items: center;
            margin-bottom: 6px;
        ">
            <span style="
                color: #06b6d4;
                font-weight: 800;
                font-size: 0.95rem;
            ">
                {status_text}
            </span>
            <span style="
                color: #f59e0b;
                font-weight: 900;
                font-size: 1.05rem;
                text-shadow: 0 0 10px rgba(245, 158, 11, 0.5);
            ">
                {percentage}%
            </span>
        </div>
        <div style="
            background: rgba(15, 23, 42, 0.9);
            border: 2px solid #10b981;
            border-radius: 12px;
            padding: 3px;
            box-shadow: 0 0 15px rgba(16, 185, 129, 0.3);
            position: relative;
            overflow: hidden;
            height: 24px;
        ">
            <div style="
                width: {percentage}%;
                height: 18px;
                background: linear-gradient(
                    90deg,
                    #10b981 0%,
                    #3b82f6 50%,
                    #f59e0b 100%
                );
                border-radius: 8px;
                box-shadow: 0 0 20px rgba(245, 158, 11, 0.8);
                transition: width 0.3s ease-in-out;
            ">
            </div>
        </div>
    </div>
    """


# =============================================================================
# 7. MULTIMODAL PAYLOAD EXTRACTION
# =============================================================================

def extract_docx_details(file):
    """
    Extract text from DOCX paragraphs, tables, headers/footers,
    and extract embedded images.
    """
    extracted_text = ""
    extracted_images = []

    try:
        file.seek(0)
        doc = docx.Document(file)

        extracted_text += f"\n--- FILE: {file.name} ---\n"

        # Headers & Footers
        for section in doc.sections:
            for header_p in section.header.paragraphs:
                if header_p.text.strip():
                    extracted_text += f"[Header] {header_p.text.strip()}\n"
            for footer_p in section.footer.paragraphs:
                if footer_p.text.strip():
                    extracted_text += f"[Footer] {footer_p.text.strip()}\n"

        # Paragraphs
        para_num = 0
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                para_num += 1
                extracted_text += f"[Paragraph {para_num}] {text}\n"

        # Tables
        for t_idx, table in enumerate(doc.tables, start=1):
            extracted_text += f"\n[TABLE {t_idx}]\n"
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                extracted_text += " | ".join(cells) + "\n"

        # Embedded Images
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                img_part = rel.target_part
                img_bytes = img_part.blob
                img_ext = img_part.content_type.split('/')[-1]
                if img_ext in ['png', 'jpeg', 'jpg']:
                    try:
                        img = Image.open(io.BytesIO(img_bytes))
                        extracted_images.append(img)
                    except Exception:
                        pass

    except Exception as exc:
        extracted_text += f"\n--- FILE: {file.name} ---\n[ERROR READING DOCUMENT: {exc}]\n"

    return extracted_text, extracted_images


def build_multimodal_payload(sow_files, notes_files, media_files, loose_notes):
    """
    Consolidates DOCX text, tables, embedded images, TXT files, loose notes, 
    and uploaded media files (png, jpg, mp3, mp4) into a multimodal payload.
    """
    payload_parts = []
    text_buffer = ""

    # 1. Process DOCX Files
    if sow_files:
        for file in sow_files:
            docx_text, embedded_imgs = extract_docx_details(file)
            text_buffer += docx_text
            payload_parts.extend(embedded_imgs)

    # 2. Process TXT Files
    if notes_files:
        for file in notes_files:
            try:
                file.seek(0)
                raw = file.read()
                text = raw.decode("utf-8", errors="replace")
                text_buffer += f"\n--- FILE: {file.name} ---\n{text}\n"
            except Exception as exc:
                text_buffer += f"\n--- FILE: {file.name} ---\n[ERROR READING TXT: {exc}]\n"

    # 3. Process Loose Notes
    if loose_notes and loose_notes.strip():
        text_buffer += f"\n--- LOOSE NOTES / CLIENT EMAILS ---\n{loose_notes.strip()}\n"

    # Insert combined text buffer as the initial text block
    if text_buffer.strip():
        payload_parts.insert(0, text_buffer)

    # 4. Process Direct Uploaded Media Files (.png, .jpg, .mp3, .mp4)
    if media_files:
        for m_file in media_files:
            try:
                m_file.seek(0)
                bytes_data = m_file.read()
                mime_type = m_file.type

                if mime_type.startswith("image/"):
                    img = Image.open(io.BytesIO(bytes_data))
                    payload_parts.append(img)
                elif mime_type.startswith("audio/") or mime_type.startswith("video/"):
                    media_part = types.Part.from_bytes(
                        data=bytes_data,
                        mime_type=mime_type
                    )
                    payload_parts.append(media_part)
            except Exception as exc:
                payload_parts.append(f"\n[ERROR PROCESSING MEDIA FILE {m_file.name}: {exc}]\n")

    return payload_parts


# =============================================================================
# 8. SMART SCOPE ANALYSIS PROMPT
# =============================================================================

SYSTEM_INSTRUCTION_PROMPT = """
You are a senior IT Delivery Lead, Business Analyst,
Product Owner and Solution Architect.

Your task is to transform messy project-intake information (documents, tables, diagrams, video, audio)
into a professional delivery-ready Smart Scope Analysis.

IMPORTANT:
Do NOT invent requirements.
Only extract requirements that are explicitly stated
or strongly implied by the supplied material.

If information is missing, classify it as a gap,
assumption, dependency or risk rather than inventing an answer.

============================================================
PRIMARY OBJECTIVE
============================================================
Produce a professional scope-handover analysis identifying:
1. What is actually requested.
2. Which module/functionality it belongs to.
3. Where the requirement came from.
4. What is missing or ambiguous.
5. Conflicts between documents.
6. Risks and dependencies.
7. Requirements suitable for Jira.
8. Testable acceptance criteria.

============================================================
QUALITY RULES
============================================================
Return ONLY valid JSON matching the exact schema provided.
No Markdown wrapping. No explanations outside JSON. Do not use trailing commas.

============================================================
JSON SCHEMA
============================================================
{
  "project_summary": {
    "project_objective": "",
    "business_goal": "",
    "overall_scope_summary": ""
  },
  "extracted_scope": [
    {
      "module": "",
      "source": "",
      "scope_type": "FUNCTIONAL/NON-FUNCTIONAL/TECHNICAL/BUSINESS",
      "points": [ "" ]
    }
  ],
  "gaps_and_risks": [
    {
      "severity": "HIGH/MEDIUM/LOW",
      "type": "",
      "description": "",
      "impact": "",
      "recommended_action": ""
    }
  ],
  "assumptions": [
    {
      "assumption": "",
      "reason": "",
      "validation_required": true
    }
  ],
  "dependencies": [
    {
      "dependency": "",
      "owner": "",
      "impact": ""
    }
  ],
  "conflicts": [
    {
      "topic": "",
      "source_a": "",
      "statement_a": "",
      "source_b": "",
      "statement_b": "",
      "resolution_needed": ""
    }
  ],
  "jira_user_stories": [
    {
      "title": "",
      "user_role": "",
      "want_statement": "",
      "so_that_statement": "",
      "acceptance_criteria": [ "" ]
    }
  ]
}
"""


# =============================================================================
# 9. GEMINI ANALYSIS
# =============================================================================

def analyze_with_gemini(multimodal_payload):

    if not client:
        st.error("GEMINI_API_KEY is missing in Streamlit Secrets.")
        return None

    # Prepend system prompt to multimodal contents payload
    contents = [SYSTEM_INSTRUCTION_PROMPT] + multimodal_payload

    try:
        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=contents,
            config=types.GenerateContentConfig(
                response_mime_type="application/json",
                temperature=0.15,
            )
        )

        if not response.text:
            st.error("Gemini returned an empty response.")
            return None

        result = json.loads(response.text)
        return result

    except json.JSONDecodeError as exc:
        st.error(f"Gemini returned invalid JSON: {exc}")
        return None
    except Exception as exc:
        st.error(f"Gemini API Error: {exc}")
        return None


# =============================================================================
# 10. MOCK ANALYSIS
# =============================================================================

MOCK_ANALYSIS = {
    "project_summary": {
        "project_objective": "Transform project intake materials into a structured delivery scope.",
        "business_goal": "Reduce manual requirement analysis and improve project handover quality.",
        "overall_scope_summary": "The solution extracts scope, identifies risks and gaps, and creates Jira-ready user stories."
    },
    "extracted_scope": [
        {
            "module": "AUTHENTICATION & SECURITY",
            "source": "Proposal Page 4, Section 2",
            "scope_type": "FUNCTIONAL",
            "points": [
                "Multi-factor authentication (MFA) via SMS or Email.",
                "Role-based access control (RBAC) for Admin and Client roles."
            ]
        },
        {
            "module": "HANDOVER INTELLIGENCE ENGINE",
            "source": "Kickoff Transcript Lines 12-45",
            "scope_type": "TECHNICAL",
            "points": [
                "Multi-format document parsing for project intake material.",
                "Automated gap auditing and structured requirement extraction.",
                "Generation of structured user stories for Jira."
            ]
        }
    ],
    "gaps_and_risks": [
        {
            "severity": "HIGH",
            "type": "Missing SLA",
            "description": "No operational SLA has been defined for document processing.",
            "impact": "Performance expectations cannot be validated during delivery.",
            "recommended_action": "Define expected processing time based on document size and complexity."
        },
        {
            "severity": "MEDIUM",
            "type": "Cross-Document Conflict",
            "description": "One project input references a 500MB upload limit while another specifies 200MB.",
            "impact": "Development and infrastructure sizing may be incorrect.",
            "recommended_action": "Confirm the authoritative maximum upload size with the client."
        }
    ],
    "assumptions": [
        {
            "assumption": "Admin and Client are the initial supported roles.",
            "reason": "These roles are explicitly mentioned in the project material.",
            "validation_required": True
        }
    ],
    "dependencies": [
        {
            "dependency": "Gemini API availability",
            "owner": "AI / Platform Team",
            "impact": "AI-powered scope extraction depends on API availability and valid credentials."
        }
    ],
    "conflicts": [
        {
            "topic": "Maximum Upload Size",
            "source_a": "Kickoff Transcript",
            "statement_a": "500MB upload support requested.",
            "source_b": "Statement of Work",
            "statement_b": "200MB upload limit specified.",
            "resolution_needed": "Confirm the authoritative maximum upload size."
        }
    ],
    "jira_user_stories": [
        {
            "title": "Configurable Document Processing Engine",
            "user_role": "Delivery Lead",
            "want_statement": "ingest project intake documents concurrently",
            "so_that_statement": "I can automatically extract structured requirements without manually reviewing every document",
            "acceptance_criteria": [
                "Given a user uploads supported intake documents",
                "When the user selects Generate Smart Scope",
                "Then the system extracts the available project requirements",
                "Given the uploaded documents contain conflicting information",
                "When the analysis is generated",
                "Then the conflicting requirements are identified for review."
            ]
        }
    ]
}


# =============================================================================
# 11. DOCX REPORT GENERATOR
# =============================================================================

def build_docx_report(data):

    doc = docx.Document()

    doc.add_heading(
        "Pitch to Project - Handover Scope Analysis",
        0
    )

    # Project Summary
    summary = data.get("project_summary", {})
    doc.add_heading("Project Summary", level=1)
    doc.add_paragraph(f"Project Objective: {summary.get('project_objective', '')}")
    doc.add_paragraph(f"Business Goal: {summary.get('business_goal', '')}")
    doc.add_paragraph(f"Overall Scope: {summary.get('overall_scope_summary', '')}")

    # Extracted Scope
    doc.add_heading("1. Extracted Scope", level=1)
    for mod in data.get("extracted_scope", []):
        doc.add_heading(f"Module: {mod.get('module', 'General')}", level=2)
        doc.add_paragraph(f"Source: {mod.get('source', 'Uploaded Documents')}")
        doc.add_paragraph(f"Scope Type: {mod.get('scope_type', 'FUNCTIONAL')}")
        for point in mod.get("points", []):
            doc.add_paragraph(f"• {point}")

    # Risks
    doc.add_heading("2. Gaps & Risk Audit", level=1)
    for gap in data.get("gaps_and_risks", []):
        doc.add_paragraph(f"[{gap.get('severity', 'INFO')}] {gap.get('type', 'Risk')}: {gap.get('description', '')}")
        if gap.get("impact"):
            doc.add_paragraph(f"Impact: {gap.get('impact')}")
        if gap.get("recommended_action"):
            doc.add_paragraph(f"Recommended Action: {gap.get('recommended_action')}")

    # Assumptions
    doc.add_heading("3. Assumptions", level=1)
    assumptions = data.get("assumptions", [])
    if assumptions:
        for assumption in assumptions:
            doc.add_paragraph(f"• {assumption.get('assumption', '')}")
            if assumption.get("reason"):
                doc.add_paragraph(f"Reason: {assumption.get('reason')}")
    else:
        doc.add_paragraph("No explicit assumptions identified.")

    # Dependencies
    doc.add_heading("4. Dependencies", level=1)
    dependencies = data.get("dependencies", [])
    if dependencies:
        for dependency in dependencies:
            doc.add_paragraph(f"• {dependency.get('dependency', '')}")
            if dependency.get("owner"):
                doc.add_paragraph(f"Owner: {dependency.get('owner')}")
            if dependency.get("impact"):
                doc.add_paragraph(f"Impact: {dependency.get('impact')}")
    else:
        doc.add_paragraph("No dependencies identified.")

    # Conflicts
    doc.add_heading("5. Requirement Conflicts", level=1)
    conflicts = data.get("conflicts", [])
    if conflicts:
        for conflict in conflicts:
            doc.add_heading(conflict.get("topic", "Conflict"), level=2)
            doc.add_paragraph(f"Source A: {conflict.get('source_a', '')}")
            doc.add_paragraph(f"Statement A: {conflict.get('statement_a', '')}")
            doc.add_paragraph(f"Source B: {conflict.get('source_b', '')}")
            doc.add_paragraph(f"Statement B: {conflict.get('statement_b', '')}")
            doc.add_paragraph(f"Resolution Needed: {conflict.get('resolution_needed', '')}")
    else:
        doc.add_paragraph("No conflicts identified.")

    # Jira Stories
    doc.add_heading("6. Jira User Stories", level=1)
    for story in data.get("jira_user_stories", []):
        doc.add_heading(story.get("title", "User Story"), level=2)
        doc.add_paragraph(
            f"As a {story.get('user_role', 'User')}, I want to {story.get('want_statement', '')} "
            f"so that {story.get('so_that_statement', '')}."
        )
        doc.add_paragraph("Acceptance Criteria:")
        for ac in story.get("acceptance_criteria", []):
            doc.add_paragraph(f"  - {ac}")

    target_stream = io.BytesIO()
    doc.save(target_stream)
    target_stream.seek(0)
    return target_stream.getvalue()


# =============================================================================
# 12. MAIN ANIMATED HEADER BANNER
# =============================================================================

components.html(
    """
    <!DOCTYPE html>
    <html>
    <head>
        <style>
            html, body { margin: 0; padding: 0; background: transparent; overflow: hidden; }
            .banner {
                width: 100%; height: 42px; border-radius: 10px; overflow: hidden; display: flex; align-items: center;
                background: linear-gradient(90deg, #ec4899 0%, #8b5cf6 33%, #06b6d4 66%, #10b981 100%);
                box-shadow: 0 0 20px rgba(139, 92, 246, 0.45);
                font-family: system-ui, -apple-system, sans-serif;
            }
            .track { display: flex; width: max-content; animation: scrollBanner 18s linear infinite; white-space: nowrap; }
            .banner-text { color: #ffffff; font-size: 25px; font-weight: 800; padding-right: 90px; text-shadow: 0 2px 6px rgba(0,0,0,0.45); }
            @keyframes scrollBanner { from { transform: translateX(0); } to { transform: translateX(-50%); } }
        </style>
    </head>
    <body>
        <div class="banner">
            <div class="track">
                <span class="banner-text">⚡ Smart Scope Handover Engine &nbsp; | &nbsp; Project Intelligence Layer</span>
                <span class="banner-text">⚡ Smart Scope Handover Engine &nbsp; | &nbsp; Project Intelligence Layer</span>
                <span class="banner-text">⚡ Smart Scope Handover Engine &nbsp; | &nbsp; Project Intelligence Layer</span>
                <span class="banner-text">⚡ Smart Scope Handover Engine &nbsp; | &nbsp; Project Intelligence Layer</span>
            </div>
        </div>
    </body>
    </html>
    """,
    height=50,
    scrolling=False
)


# =============================================================================
# 13. APPLICATION TITLE
# =============================================================================

col_title, col_toggle = st.columns([3, 1])

with col_title:
    components.html(
        """
        <div style="font-family: system-ui, -apple-system, sans-serif;">
            <div style="color: #ffffff; font-size: 2.8rem; font-weight: 900; margin: 0; line-height: 1.2; text-shadow: 0 0 15px rgba(56, 189, 248, 0.6);">
                Pitch to Project
            </div>
            <div style="color: #38bdf8; font-size: 1.05rem; font-weight: 700; margin-top: 6px; line-height: 1.3;">
                🚀 AI-Powered Scope Intelligence & Handover Engine
            </div>
        </div>
        """,
        height=120,
        scrolling=False
    )

with col_toggle:
    demo_mode = st.toggle("Demo Mode (Safe Pitch)", value=True)


# =============================================================================
# 14. MAIN TWO-COLUMN LAYOUT
# =============================================================================

left_col, right_col = st.columns([1, 1], gap="medium")


# =============================================================================
# 15. LEFT COLUMN - INPUT
# =============================================================================

with left_col:

    st.header("1. Intake Documents & Media")

    sow_files = st.file_uploader(
        "Proposals / SOWs (.docx)",
        type=["docx"],
        accept_multiple_files=True
    )

    notes_files = st.file_uploader(
        "Transcripts / Notes (.txt)",
        type=["txt"],
        accept_multiple_files=True
    )

    media_files = st.file_uploader(
        "Diagrams & Media (.png, .jpg, .mp4, .mp3)",
        type=["png", "jpg", "mp4", "mp3"],
        accept_multiple_files=True
    )

    loose_notes = st.text_area(
        "Or Paste Loose Client Emails / Notes:",
        height=120,
        placeholder="Paste kickoff notes or client requirements here..."
    )

    generate_btn = st.button("⚡ GENERATE SMART SCOPE")


# =============================================================================
# 16. RIGHT COLUMN - AI ANALYSIS (CONTAINER PERSISTENCE FIXED)
# =============================================================================

with right_col:

    st.header("2. AI Scope & Handover Analysis")

    if generate_btn:

        progress_card = st.empty()

        # DEMO MODE
        if demo_mode:
            with progress_card.container(border=True):
                st.markdown("### 🧠 Running Scope Analysis")
                bar_ph = st.empty()

                for i in range(101):
                    time.sleep(0.01)
                    bar_ph.markdown(
                        render_stylish_progress(i, "⚙️ Processing intake materials & generating user stories..."),
                        unsafe_allow_html=True
                    )

                st.success("✅ Demo Analysis Loaded Successfully!")
                time.sleep(0.4)

            progress_card.empty()
            st.session_state["analysis_data"] = MOCK_ANALYSIS
            st.toast("⚡ Demo Analysis Loaded!", icon="✅")

        # LIVE GEMINI MODE
        else:
            with progress_card.container(border=True):
                st.markdown("### 🧠 Live Gemini Multimodal Scope Extraction")
                bar_ph = st.empty()

                # Step 1: Processing
                bar_ph.markdown(
                    render_stylish_progress(20, "📄 Step 1/3: Extracting text, tables & embedded doc images..."),
                    unsafe_allow_html=True
                )
                
                payload = build_multimodal_payload(
                    sow_files,
                    notes_files,
                    media_files,
                    loose_notes
                )

                if not payload:
                    st.warning("Please upload at least one document/media file or paste notes before analyzing.")
                    time.sleep(2)
                    progress_card.empty()
                else:
                    # Step 2: Transmitting
                    bar_ph.markdown(
                        render_stylish_progress(50, "⚡ Step 2/3: Transmitting multimodal payload to Gemini..."),
                        unsafe_allow_html=True
                    )
                    time.sleep(0.4)

                    # Step 3: Auditing
                    bar_ph.markdown(
                        render_stylish_progress(80, "🔍 Step 3/3: Auditing scope, conflicts, media & Jira stories..."),
                        unsafe_allow_html=True
                    )

                    result = analyze_with_gemini(payload)

                    if result:
                        bar_ph.markdown(
                            render_stylish_progress(100, "✅ Smart Scope Handover Analysis Complete!"),
                            unsafe_allow_html=True
                        )
                        st.session_state["analysis_data"] = result
                        
                        # Retain completed 100% status state visible before clearing UI frame
                        time.sleep(1.5)
                        progress_card.empty()
                        
                        st.toast("⚡ Live Gemini Extraction Complete!", icon="✅")
                        st.rerun()
                    else:
                        st.error("❌ Extraction Failed")
                        time.sleep(2)
                        progress_card.empty()

    # DEFAULT DATA
    if "analysis_data" not in st.session_state:
        st.session_state["analysis_data"] = MOCK_ANALYSIS

    data = st.session_state["analysis_data"]

    # TABS
    tab_summary, tab_scope, tab_risks, tab_jira = st.tabs([
        "📊 Project Summary",
        "📌 Extracted Scope",
        "🚨 Missing Items & Risks",
        "🚀 Jira User Stories"
    ])

    # TAB 1: SUMMARY
    with tab_summary:
        summary = data.get("project_summary", {})

        with st.container(border=True):
            st.markdown("### 🎯 Project Objective")
            st.write(summary.get("project_objective", "Not provided."))

        with st.container(border=True):
            st.markdown("### 💼 Business Goal")
            st.write(summary.get("business_goal", "Not provided."))

        with st.container(border=True):
            st.markdown("### 📋 Overall Scope")
            st.write(summary.get("overall_scope_summary", "Not provided."))

        st.markdown("### 🧠 Assumptions")
        assumptions = data.get("assumptions", [])
        if assumptions:
            for item in assumptions:
                with st.container(border=True):
                    st.markdown(f"**{item.get('assumption', '')}**")
                    if item.get("reason"):
                        st.caption(f"Reason: {item.get('reason')}")
        else:
            st.info("No assumptions identified.")

    # TAB 2: EXTRACTED SCOPE
    with tab_scope:
        scope_items = data.get("extracted_scope", [])
        if not scope_items:
            st.info("No scope items were identified.")
        for item in scope_items:
            with st.container(border=True):
                st.markdown(f"### 📌 MODULE: {item.get('module', 'General Scope')}")
                st.caption(f"Source: {item.get('source', 'Uploaded Files')}")
                st.caption(f"Type: {item.get('scope_type', 'FUNCTIONAL')}")
                for pt in item.get("points", []):
                    st.markdown(f"• {pt}")

    # TAB 3: RISKS & CONFLICTS
    with tab_risks:
        risks = data.get("gaps_and_risks", [])
        if not risks:
            st.success("✅ No significant gaps or risks identified.")
        for risk in risks:
            with st.container(border=True):
                severity = risk.get("severity", "INFO").upper()
                badge = "🔴" if severity == "HIGH" else "🟡" if severity == "MEDIUM" else "🔵"
                st.markdown(f"### {badge} [{severity}] {risk.get('type', 'Risk')}")
                st.write(risk.get("description", ""))
                if risk.get("impact"):
                    st.markdown(f"**Impact:** {risk.get('impact')}")
                if risk.get("recommended_action"):
                    st.markdown(f"**Recommended Action:** {risk.get('recommended_action')}")

        st.markdown("### 🔗 Dependencies")
        dependencies = data.get("dependencies", [])
        if dependencies:
            for dep in dependencies:
                with st.container(border=True):
                    st.markdown(f"**{dep.get('dependency', '')}**")
                    if dep.get("owner"):
                        st.caption(f"Owner: {dep.get('owner')}")
                    if dep.get("impact"):
                        st.write(dep.get("impact"))
        else:
            st.info("No dependencies identified.")

        st.markdown("### ⚔️ Requirement Conflicts")
        conflicts = data.get("conflicts", [])
        if conflicts:
            for conflict in conflicts:
                with st.container(border=True):
                    st.markdown(f"### ⚠️ {conflict.get('topic', 'Conflict')}")
                    st.markdown(f"**Source A:** {conflict.get('source_a', '')}")
                    st.write(conflict.get("statement_a", ""))
                    st.markdown(f"**Source B:** {conflict.get('source_b', '')}")
                    st.write(conflict.get("statement_b", ""))
                    st.markdown(f"**Resolution Needed:** {conflict.get('resolution_needed', '')}")
        else:
            st.success("✅ No requirement conflicts identified.")

    # TAB 4: JIRA STORIES
    with tab_jira:
        stories = data.get("jira_user_stories", [])
        if not stories:
            st.info("No Jira user stories were generated.")
        for story in stories:
            with st.container(border=True):
                st.markdown(f"### 🚀 {story.get('title', 'User Story')}")
                st.markdown(
                    f"**As a** `{story.get('user_role', 'User')}`, "
                    f"**I want to** {story.get('want_statement', '')} "
                    f"**so that** {story.get('so_that_statement', '')}."
                )
                st.markdown("**Acceptance Criteria:**")
                for ac in story.get("acceptance_criteria", []):
                    st.markdown(f"- `{ac}`")

    # EXPORT
    st.divider()
    docx_bytes = build_docx_report(data)
    st.download_button(
        label="📄 Export Handover Report (.docx)",
        data=docx_bytes,
        file_name="Pitch_to_Project_Handover_Report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )
