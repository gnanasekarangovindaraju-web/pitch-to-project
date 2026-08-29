import json
import time
import io
import mimetypes

import docx
import streamlit as st
import streamlit.components.v1 as components

from google import genai
from google.genai import types


# =============================================================================
# 1. PAGE CONFIGURATION
# =============================================================================

st.set_page_config(
    page_title="Pitch to Project | Smart Scope Engine",
    page_icon="⚡",
    layout="wide",
)


# =============================================================================
# 2. GLOBAL CSS
#    Existing UI design preserved
# =============================================================================

css_code = """
<style>

/* ============================================================
   GLOBAL BACKGROUND
   ============================================================ */

.stApp {
    background: linear-gradient(
        125deg,
        #0f172a 0%,
        #1e1b4b 35%,
        #311042 70%,
        #0284c7 100%
    ) !important;

    background-attachment: fixed;
}


/* ============================================================
   LOGIN FORM
   ============================================================ */

div[data-testid="stForm"] {
    background: rgba(15, 23, 42, 0.95) !important;
    border: 2.5px solid #a855f7 !important;
    border-radius: 18px !important;
    padding: 36px !important;
    box-shadow: 0 0 40px rgba(168, 85, 247, 0.5) !important;
}


/* ============================================================
   LABELS
   ============================================================ */

div[data-testid="stForm"] label,
div[data-testid="stForm"] label p,
div[data-testid="stTextInput"] label p {
    color: #38bdf8 !important;
    font-weight: 900 !important;
    font-size: 1.4rem !important;
    letter-spacing: 0.5px !important;
    margin-bottom: 8px !important;
    text-align: left !important;
    width: 100% !important;
    display: block !important;
}


/* ============================================================
   INPUTS
   ============================================================ */

div[data-testid="stForm"] div[data-testid="stTextInput"] input,
div[data-testid="stTextInput"] input,
input[type="text"],
input[type="password"] {

    background-color: #0f172a !important;
    color: #ffffff !important;
    -webkit-text-fill-color: #ffffff !important;

    border: 2.5px solid #38bdf8 !important;
    border-radius: 12px !important;

    font-weight: 800 !important;
    font-size: 1.35rem !important;

    padding: 16px 20px !important;

    text-align: left !important;

    box-shadow:
        0 0 14px rgba(56, 189, 248, 0.3) !important;
}


/* ============================================================
   AUTOFILL
   ============================================================ */

input:-webkit-autofill,
input:-webkit-autofill:hover,
input:-webkit-autofill:focus,
input:-webkit-autofill:active {

    -webkit-box-shadow:
        0 0 0 1000px #0f172a inset !important;

    -webkit-text-fill-color: #ffffff !important;

    caret-color: #ffffff !important;

    border: 2.5px solid #38bdf8 !important;
}


/* ============================================================
   PLACEHOLDER
   ============================================================ */

div[data-testid="stTextInput"] input::placeholder,
input::placeholder {

    color: #cbd5e1 !important;
    -webkit-text-fill-color: #cbd5e1 !important;

    font-weight: 700 !important;
    font-size: 1.25rem !important;

    opacity: 1 !important;

    text-align: left !important;
}


/* ============================================================
   PASSWORD ICON
   ============================================================ */

div[data-testid="stTextInput"] button svg,
div[data-testid="stForm"] svg {

    fill: #38bdf8 !important;
    stroke: #38bdf8 !important;

    width: 26px !important;
    height: 26px !important;
}


/* ============================================================
   LOGIN BUTTON
   ============================================================ */

div[data-testid="stForm"] button[type="submit"],
div[data-testid="stForm"] button[data-testid="stFormSubmitButton"],
div[data-testid="stForm"] button {

    background:
        linear-gradient(
            90deg,
            #ec4899 0%,
            #8b5cf6 50%,
            #06b6d4 100%
        ) !important;

    border: none !important;

    border-radius: 14px !important;

    padding: 18px 32px !important;

    box-shadow:
        0 0 30px rgba(236, 72, 153, 0.7) !important;

    transition: all 0.3s ease !important;

    margin-top: 22px !important;

    width: 100% !important;
}


div[data-testid="stForm"] button[type="submit"] *,
div[data-testid="stForm"] button[type="submit"] p,
div[data-testid="stForm"] button[type="submit"] span {

    color: #ffffff !important;
    -webkit-text-fill-color: #ffffff !important;

    font-weight: 900 !important;
    font-size: 1.5rem !important;

    letter-spacing: 1.2px !important;
}


/* ============================================================
   SIDEBAR
   ============================================================ */

section[data-testid="stSidebar"] {

    background: rgba(15, 23, 42, 0.95) !important;

    border-right:
        1.5px solid #a855f7 !important;
}


section[data-testid="stSidebar"] h3 {

    color: #38bdf8 !important;

    font-size: 1.3rem !important;

    font-weight: 800 !important;
}


section[data-testid="stSidebar"] p,
section[data-testid="stSidebar"] span,
section[data-testid="stSidebar"] div {

    color: #f8fafc !important;

    font-size: 1.05rem !important;

    font-weight: 700 !important;

    background: transparent !important;
}


section[data-testid="stSidebar"] div.stButton > button {

    background:
        linear-gradient(
            90deg,
            #ec4899 0%,
            #f43f5e 100%
        ) !important;

    color: #ffffff !important;

    font-weight: 800 !important;

    font-size: 1.1rem !important;

    border-radius: 10px !important;

    box-shadow:
        0 0 15px rgba(244, 63, 94, 0.5) !important;

    border: none !important;

    margin-top: 10px !important;
}


/* ============================================================
   MAIN HEADINGS
   ============================================================ */

h2,
h3 {

    color: #f8fafc !important;

    font-weight: 800 !important;

    text-shadow:
        0 0 10px rgba(168, 85, 247, 0.3) !important;
}


/* ============================================================
   GENERAL TEXT
   ============================================================ */

div[data-testid="stMarkdownContainer"] p,
label[data-testid="stWidgetLabel"] p,
div[data-testid="stToggle"] span {

    color: #f8fafc !important;

    font-weight: 700 !important;
}


/* ============================================================
   FILE UPLOADER
   ============================================================ */

div[data-testid="stFileUploader"] {

    background:
        rgba(15, 23, 42, 0.95) !important;

    border:
        2px solid #a855f7 !important;

    border-radius:
        14px !important;

    padding:
        12px !important;
}


div[data-testid="stFileUploader"] section,
div[data-testid="stFileUploaderDropzone"],
div[data-testid="stFileUploader"]
[data-testid="stFileUploaderDropzone"] {

    background:
        #1e1b4b !important;

    border:
        2px dashed #38bdf8 !important;

    border-radius:
        10px !important;
}


div[data-testid="stFileUploaderDropzone"] *,
div[data-testid="stFileUploaderDropzone"] span,
div[data-testid="stFileUploaderDropzone"] small,
div[data-testid="stFileUploaderDropzone"] p {

    color: #ffffff !important;

    font-weight: 700 !important;
}


/* ============================================================
   TOAST
   ============================================================ */

div[data-testid="stToast"],
div[data-testid="stToast"] > div {

    background-color:
        #1e1b4b !important;

    background:
        #1e1b4b !important;

    border:
        2px solid #38bdf8 !important;

    border-radius:
        12px !important;

    box-shadow:
        0 0 20px rgba(56, 189, 248, 0.5) !important;
}


div[data-testid="stToast"] * {

    color: #ffffff !important;

    font-weight: 800 !important;
}


/* ============================================================
   DOWNLOAD BUTTON
   ============================================================ */

div[data-testid="stDownloadButton"] > button {

    background:
        linear-gradient(
            135deg,
            #10b981 0%,
            #059669 100%
        ) !important;

    border-radius:
        12px !important;

    border:
        none !important;

    box-shadow:
        0 0 18px rgba(16, 185, 129, 0.5) !important;

    width:
        100%;
}


div[data-testid="stDownloadButton"] > button * {

    color: #ffffff !important;

    font-weight: 900 !important;
}


/* ============================================================
   TEXT AREA
   ============================================================ */

div[data-testid="stTextArea"] textarea {

    background:
        rgba(15, 23, 42, 0.85) !important;

    border:
        2px solid #a855f7 !important;

    border-radius:
        14px !important;

    color:
        #ffffff !important;

    -webkit-text-fill-color:
        #ffffff !important;
}


div[data-testid="stTextArea"] textarea::placeholder {

    color:
        #cbd5e1 !important;

    -webkit-text-fill-color:
        #cbd5e1 !important;

    opacity:
        1 !important;
}


/* ============================================================
   BUTTONS
   ============================================================ */

div.stButton > button {

    background:
        linear-gradient(
            90deg,
            #ec4899 0%,
            #8b5cf6 50%,
            #3b82f6 100%
        ) !important;

    color:
        #ffffff !important;

    font-weight:
        800 !important;

    font-size:
        1.05rem !important;

    border-radius:
        12px !important;

    border:
        none !important;

    padding:
        14px 28px !important;

    box-shadow:
        0 0 20px rgba(139, 92, 246, 0.5) !important;

    width:
        100%;
}


button[aria-selected="true"] {

    background:
        linear-gradient(
            135deg,
            #8b5cf6 0%,
            #ec4899 100%
        ) !important;

    color:
        #ffffff !important;
}


/* ============================================================
   ANALYSIS CARDS
   ============================================================ */

div[data-testid="stVerticalBlockBorderWrapper"] > div {

    background:
        rgba(15, 23, 42, 0.8) !important;

    backdrop-filter:
        blur(10px) !important;

    border-left:
        6px solid #06b6d4 !important;

    border-radius:
        14px !important;

    padding:
        20px !important;
}


/* ============================================================
   CODE
   ============================================================ */

code {

    background-color:
        rgba(30, 27, 75, 0.95) !important;

    color:
        #38bdf8 !important;

    border:
        1px solid #a855f7 !important;

    border-radius:
        6px !important;

    padding:
        3px 8px !important;
}

</style>
"""

st.markdown(css_code, unsafe_allow_html=True)


# =============================================================================
# 3. AUTHENTICATION
# =============================================================================

def check_password():

    if st.session_state.get("authenticated", False):
        return True

    st.markdown("<br><br>", unsafe_allow_html=True)

    col1, col2, col3 = st.columns([1, 2.4, 1])

    with col2:

        with st.form("login_form"):

            components.html(
                """
                <div style="
                    background:
                        linear-gradient(
                            135deg,
                            #ec4899 0%,
                            #8b5cf6 50%,
                            #06b6d4 100%
                        );

                    border-radius: 14px;

                    padding: 24px 10px;

                    box-shadow:
                        0 0 25px rgba(236, 72, 153, 0.6);

                    text-align: center;

                    font-family:
                        system-ui,
                        -apple-system,
                        sans-serif;
                ">

                    <div style="
                        color: #ffffff;
                        font-size: 2.8rem;
                        font-weight: 900;
                        margin-bottom: 6px;
                        text-shadow:
                            0 3px 12px rgba(0,0,0,0.8);
                        letter-spacing: -0.5px;
                    ">
                        🔒 Pitch to Project
                    </div>

                    <div style="
                        color: #ffffff;
                        font-size: 1.35rem;
                        font-weight: 800;
                        letter-spacing: 1px;
                        text-shadow:
                            0 2px 8px rgba(0,0,0,0.8);
                    ">
                        ⚡ Scope Intelligence Engine Access
                    </div>

                </div>
                """,
                height=140
            )

            username = st.text_input(
                "Username",
                placeholder="Enter username"
            )

            password = st.text_input(
                "Password",
                type="password",
                placeholder="Enter password"
            )

            submit = st.form_submit_button(
                "🔑 LOGIN TO ENGINE"
            )

            if submit:

                valid_user = st.secrets.get(
                    "APP_USER",
                    "admin"
                )

                valid_password = st.secrets.get(
                    "APP_PASSWORD",
                    "project2026"
                )

                if (
                    username == valid_user
                    and password == valid_password
                ):

                    st.session_state["authenticated"] = True

                    st.toast(
                        "⚡ Login Successful!",
                        icon="✅"
                    )

                    st.rerun()

                else:

                    st.error(
                        "❌ Invalid Username or Password"
                    )

    return False


if not check_password():
    st.stop()


# =============================================================================
# 4. SIDEBAR
# =============================================================================

with st.sidebar:

    st.markdown("### 👤 User Session")

    st.write(
        "Logged in as **Admin**"
    )

    if st.button(
        "🚪 Logout",
        use_container_width=True
    ):

        st.session_state["authenticated"] = False

        st.rerun()


# =============================================================================
# 5. GEMINI CLIENT
# =============================================================================

@st.cache_resource
def get_gemini_client():

    try:

        api_key = st.secrets["GEMINI_API_KEY"]

        return genai.Client(
            api_key=api_key
        )

    except Exception:

        return None


client = get_gemini_client()


# =============================================================================
# 6. MODEL CONFIGURATION
# =============================================================================

GEMINI_MODEL = st.secrets.get(
    "GEMINI_MODEL",
    "gemini-3.7-flash"
)


# =============================================================================
# 7. PROGRESS BAR
# =============================================================================

def render_stylish_progress(
    percentage,
    status_text
):

    return f"""
    <div style="margin: 10px 0 18px 0;">

        <div style="
            display:flex;
            justify-content:space-between;
            align-items:center;
            margin-bottom:6px;
        ">

            <span style="
                color:#06b6d4;
                font-weight:800;
                font-size:0.95rem;
            ">
                {status_text}
            </span>

            <span style="
                color:#f59e0b;
                font-weight:900;
                font-size:1.05rem;
                text-shadow:
                    0 0 10px rgba(245,158,11,0.5);
            ">
                {percentage}%
            </span>

        </div>

        <div style="
            background:rgba(15,23,42,0.9);
            border:2px solid #10b981;
            border-radius:12px;
            padding:3px;
            box-shadow:
                0 0 15px rgba(16,185,129,0.3);
            position:relative;
            overflow:hidden;
        ">

            <div style="
                width:{percentage}%;
                height:18px;
                background:
                    linear-gradient(
                        90deg,
                        #10b981 0%,
                        #3b82f6 50%,
                        #f59e0b 100%
                    );
                border-radius:8px;
                box-shadow:
                    0 0 20px rgba(245,158,11,0.8);
                transition:
                    width 0.2s ease-in-out;
            ">
            </div>

        </div>

    </div>
    """


# =============================================================================
# 8. DOCX EXTRACTION
# =============================================================================

def extract_docx_content(file):

    """
    Extract:
        - paragraphs
        - tables

    This is intentionally done locally so the application
    has reliable DOCX processing independent of Gemini's
    direct DOCX input support.
    """

    output = []

    try:

        doc = docx.Document(file)

        output.append(
            f"\n===== DOCUMENT: {file.name} =====\n"
        )

        # ---------------------------------------------------------
        # Paragraphs
        # ---------------------------------------------------------

        paragraph_number = 0

        for paragraph in doc.paragraphs:

            text = paragraph.text.strip()

            if not text:
                continue

            paragraph_number += 1

            output.append(
                f"[{file.name} | Paragraph {paragraph_number}] "
                f"{text}"
            )


        # ---------------------------------------------------------
        # Tables
        # ---------------------------------------------------------

        for table_index, table in enumerate(
            doc.tables,
            start=1
        ):

            output.append(
                f"\n[{file.name} | TABLE {table_index}]\n"
            )

            for row_index, row in enumerate(
                table.rows,
                start=1
            ):

                cells = []

                for cell in row.cells:

                    cell_text = cell.text.strip()

                    cells.append(
                        cell_text
                    )

                output.append(
                    f"Row {row_index}: "
                    + " | ".join(cells)
                )


        return "\n".join(output)

    except Exception as exc:

        return (
            f"\n===== DOCUMENT: {file.name} =====\n"
            f"[ERROR READING DOCX] {str(exc)}\n"
        )


# =============================================================================
# 9. TEXT EXTRACTION
# =============================================================================

def extract_text_from_uploads(
    sow_files,
    notes_files,
    loose_notes
):

    combined_text = ""

    # -------------------------------------------------------------
    # DOCX
    # -------------------------------------------------------------

    if sow_files:

        for file in sow_files:

            combined_text += (
                extract_docx_content(file)
                + "\n"
            )


    # -------------------------------------------------------------
    # TXT
    # -------------------------------------------------------------

    if notes_files:

        for file in notes_files:

            try:

                file.seek(0)

                text = file.read().decode(
                    "utf-8",
                    errors="replace"
                )

                combined_text += (
                    f"\n===== FILE: {file.name} =====\n"
                )

                combined_text += text

                combined_text += "\n"

            except Exception as exc:

                combined_text += (
                    f"\n===== FILE: {file.name} =====\n"
                    f"[ERROR READING FILE] {str(exc)}\n"
                )


    # -------------------------------------------------------------
    # Pasted notes
    # -------------------------------------------------------------

    if loose_notes and loose_notes.strip():

        combined_text += (
            "\n===== LOOSE CLIENT EMAILS / NOTES =====\n"
        )

        combined_text += (
            loose_notes.strip()
        )

        combined_text += "\n"


    return combined_text


# =============================================================================
# 10. GEMINI FILE UPLOAD
# =============================================================================

def upload_media_to_gemini(
    media_files
):

    """
    Upload images/audio/video through Gemini Files API.

    Returns:
        list of uploaded Gemini File objects
    """

    uploaded_files = []

    if not media_files:
        return uploaded_files


    if not client:
        raise RuntimeError(
            "Gemini client is not available."
        )


    for uploaded in media_files:

        # ---------------------------------------------------------
        # Convert Streamlit UploadedFile into temporary file
        # ---------------------------------------------------------

        suffix = ""

        if uploaded.name.lower().endswith(".png"):
            suffix = ".png"

        elif uploaded.name.lower().endswith(
            (".jpg", ".jpeg")
        ):
            suffix = ".jpg"

        elif uploaded.name.lower().endswith(".mp4"):
            suffix = ".mp4"

        elif uploaded.name.lower().endswith(".mp3"):
            suffix = ".mp3"


        import tempfile

        with tempfile.NamedTemporaryFile(
            delete=False,
            suffix=suffix
        ) as temp_file:

            temp_file.write(
                uploaded.getvalue()
            )

            temp_path = temp_file.name


        # ---------------------------------------------------------
        # MIME type
        # ---------------------------------------------------------

        mime_type = (
            uploaded.type
            or mimetypes.guess_type(
                uploaded.name
            )[0]
        )


        # ---------------------------------------------------------
        # Upload
        # ---------------------------------------------------------

        try:

            uploaded_gemini_file = (
                client.files.upload(
                    file=temp_path
                )
            )

            uploaded_files.append(
                uploaded_gemini_file
            )

        except Exception as exc:

            raise RuntimeError(
                f"Failed to upload "
                f"{uploaded.name}: {exc}"
            )


    return uploaded_files


# =============================================================================
# 11. WAIT FOR MEDIA PROCESSING
# =============================================================================

def wait_for_gemini_files(
    uploaded_files,
    status_callback=None
):

    """
    Videos can remain in PROCESSING state for some time.
    Wait until Gemini reports ACTIVE.
    """

    ready_files = []

    total = len(uploaded_files)

    for index, gemini_file in enumerate(
        uploaded_files,
        start=1
    ):

        current_file = gemini_file

        while True:

            state = getattr(
                current_file,
                "state",
                None
            )

            state_name = getattr(
                state,
                "name",
                str(state)
            )

            if status_callback:

                percentage = int(
                    ((index - 1) / max(total, 1))
                    * 100
                )

                status_callback(
                    percentage,
                    (
                        f"🎞️ Processing media "
                        f"{index}/{total}: "
                        f"{getattr(current_file, 'display_name', 'media')}"
                    )
                )


            if state_name == "ACTIVE":

                ready_files.append(
                    current_file
                )

                break


            if state_name in (
                "FAILED",
                "ERROR"
            ):

                error_info = getattr(
                    current_file,
                    "error",
                    None
                )

                raise RuntimeError(
                    "Gemini media processing failed: "
                    f"{error_info}"
                )


            time.sleep(2)

            current_file = client.files.get(
                name=current_file.name
            )


    return ready_files


# =============================================================================
# 12. MASTER AI PROMPT
# =============================================================================

def build_scope_prompt(
    raw_text,
    media_files
):

    media_names = []

    for media in media_files:

        media_names.append(
            getattr(
                media,
                "display_name",
                "Uploaded Media"
            )
        )


    media_context = (
        ", ".join(media_names)
        if media_names
        else "No media files uploaded."
    )


    prompt = f"""
You are the Lead Business Analyst, Product Manager,
Solution Architect, QA Lead, and IT Delivery Lead working
together as a single expert.

Your job is to transform incomplete project-intake material
into a reliable Project Handover Scope Analysis.

The intake can contain:

- Proposals
- Statements of Work
- Client requirements
- Emails
- Meeting notes
- Kickoff transcripts
- Tables
- Diagrams
- Screenshots
- Audio recordings
- Videos
- Architecture discussions
- Business requirements
- Technical requirements

IMPORTANT RULES
===============

1. DO NOT invent requirements.

2. Do not assume that something is required merely because
   it is common in similar projects.

3. Clearly distinguish:
   - Explicit requirement
   - Implied requirement
   - Assumption
   - Missing information
   - Risk
   - Contradiction

4. When two documents conflict, DO NOT choose one silently.
   Report the conflict.

5. Preserve important numbers exactly:
   - dates
   - amounts
   - percentages
   - SLAs
   - response times
   - file sizes
   - user counts
   - limits
   - frequencies
   - deadlines

6. Treat diagrams and screenshots as evidence.
   Inspect visible:
   - components
   - labels
   - workflows
   - arrows
   - screens
   - entities
   - integrations
   - architecture elements
   - data flows

7. For audio/video:
   identify relevant business discussions,
   decisions, requirements, constraints,
   action items, and unresolved questions.

8. Do not convert a suggestion into a confirmed requirement.

9. If the source says something is optional,
   preserve it as optional.

10. If the source is ambiguous, mark it as ambiguous.

11. Every extracted requirement should contain a source
    whenever possible.

12. User stories must be implementation-ready enough
    for Jira refinement, but must remain faithful to
    the intake material.

13. Acceptance criteria should be testable.

14. Identify missing information that could cause:
    - scope creep
    - estimation problems
    - development blockers
    - QA ambiguity
    - production incidents
    - contractual disputes

15. Prioritize HIGH risks that can materially affect:
    - cost
    - schedule
    - architecture
    - security
    - compliance
    - customer acceptance
    - production operations

SOURCE MEDIA FILES
==================

The following media files are available for analysis:

{media_context}


TEXTUAL INTAKE
==============

{raw_text}


ANALYSIS OBJECTIVE
==================

Produce a complete handover analysis covering:

A. PROJECT SCOPE
----------------

Extract the actual project scope.

Group related requirements into logical modules.

Examples:

- Authentication & Security
- User Management
- Customer Management
- Dashboard
- Reporting
- Payments
- Notifications
- Integrations
- Data Migration
- Administration
- Infrastructure
- Handover / Operations

Do NOT force requirements into these examples.
Create modules based on the actual source material.


B. FUNCTIONAL REQUIREMENTS
---------------------------

Identify what the system must DO.

Each requirement should be concise but specific.


C. NON-FUNCTIONAL REQUIREMENTS
------------------------------

Look for:

- Performance
- Availability
- Scalability
- Security
- Accessibility
- Usability
- Auditability
- Logging
- Monitoring
- Backup
- Recovery
- Compliance
- Browser/device support
- Data retention


D. DEPENDENCIES
---------------

Identify:

- Third-party systems
- APIs
- Vendors
- Client dependencies
- Infrastructure dependencies
- Data dependencies
- Authentication providers
- External services


E. ASSUMPTIONS
--------------

Only include assumptions that are genuinely necessary
to interpret the intake.

Mark them clearly as assumptions.


F. OUT OF SCOPE
---------------

Identify items explicitly stated as excluded.

Do not invent exclusions.


G. GAPS AND RISKS
-----------------

Look specifically for:

- Missing acceptance criteria
- Missing SLA
- Missing security requirements
- Missing user roles
- Missing data ownership
- Missing API specifications
- Missing integration details
- Missing error handling
- Missing reporting requirements
- Missing deployment requirements
- Missing environment details
- Missing backup/recovery requirements
- Missing monitoring requirements
- Missing support model
- Missing ownership
- Missing deadlines
- Missing volume estimates
- Conflicting requirements
- Ambiguous requirements


H. CLIENT QUESTIONS
-------------------

Generate practical clarification questions.

Questions should help the delivery team resolve
ambiguity BEFORE development begins.


I. JIRA USER STORIES
--------------------

Generate user stories only for meaningful functional requirements.

Use:

As a <role>,
I want to <action>,
so that <business value>.

Acceptance criteria should use:

Given ...
When ...
Then ...

Do not create duplicate stories.


J. HANDOVER SUMMARY
-------------------

Provide:

- Executive summary
- Overall scope confidence
- Major delivery concerns
- Recommended next actions


JSON OUTPUT
===========

Return ONLY valid JSON.

Use exactly this top-level structure:

{{
  "project_summary": {{
    "project_name": "Unknown if not provided",
    "executive_summary": "",
    "scope_confidence": "HIGH/MEDIUM/LOW",
    "major_concerns": [],
    "recommended_next_actions": []
  }},

  "extracted_scope": [
    {{
      "module": "",
      "source": "",
      "confidence": "HIGH/MEDIUM/LOW",
      "requirements": [
        {{
          "requirement_id": "REQ-001",
          "type": "FUNCTIONAL/NON_FUNCTIONAL",
          "description": "",
          "source": "",
          "status": "CONFIRMED/IMPLIED/AMBIGUOUS"
        }}
      ]
    }}
  ],

  "dependencies": [
    {{
      "dependency": "",
      "description": "",
      "source": "",
      "impact": "HIGH/MEDIUM/LOW"
    }}
  ],

  "assumptions": [
    {{
      "assumption": "",
      "reason": "",
      "impact": "HIGH/MEDIUM/LOW"
    }}
  ],

  "out_of_scope": [
    {{
      "item": "",
      "source": ""
    }}
  ],

  "gaps_and_risks": [
    {{
      "severity": "HIGH/MEDIUM/LOW",
      "type": "",
      "description": "",
      "impact": "",
      "source": "",
      "recommended_action": ""
    }}
  ],

  "client_questions": [
    {{
      "priority": "HIGH/MEDIUM/LOW",
      "question": "",
      "reason": ""
    }}
  ],

  "jira_user_stories": [
    {{
      "story_id": "US-001",
      "title": "",
      "user_role": "",
      "want_statement": "",
      "so_that_statement": "",
      "source": "",
      "acceptance_criteria": []
    }}
  ]
}}

FINAL VALIDATION
================

Before returning JSON:

- Ensure valid JSON.
- Ensure every top-level key exists.
- Ensure arrays are arrays.
- Do not return Markdown.
- Do not wrap JSON inside ```json.
- Do not add commentary outside JSON.
- Do not invent source references.
- Do not invent requirements.
"""

    return prompt


# =============================================================================
# 13. GEMINI ANALYSIS
# =============================================================================

def analyze_with_gemini(
    raw_text,
    uploaded_media_files
):

    if not client:

        st.error(
            "GEMINI_API_KEY is missing "
            "in Streamlit Secrets."
        )

        return None


    prompt = build_scope_prompt(
        raw_text,
        uploaded_media_files
    )


    contents = []

    # -------------------------------------------------------------
    # Prompt
    # -------------------------------------------------------------

    contents.append(prompt)


    # -------------------------------------------------------------
    # Media
    # -------------------------------------------------------------

    for media in uploaded_media_files:

        contents.append(
            media
        )


    # -------------------------------------------------------------
    # Gemini request
    # -------------------------------------------------------------

    try:

        response = client.models.generate_content(

            model=GEMINI_MODEL,

            contents=contents,

            config=types.GenerateContentConfig(

                response_mime_type="application/json",

                temperature=0.1,

            )
        )


        if not response.text:

            st.error(
                "Gemini returned an empty response."
            )

            return None


        # ---------------------------------------------------------
        # Parse JSON
        # ---------------------------------------------------------

        result = json.loads(
            response.text
        )


        return result


    except json.JSONDecodeError as exc:

        st.error(
            "Gemini returned invalid JSON: "
            f"{exc}"
        )

        return None


    except Exception as exc:

        st.error(
            "Gemini API Error: "
            f"{str(exc)}"
        )

        return None


# =============================================================================
# 14. DOCX REPORT
# =============================================================================

def build_docx_report(data):

    doc = docx.Document()


    # -------------------------------------------------------------
    # Title
    # -------------------------------------------------------------

    doc.add_heading(
        "Pitch to Project - Handover Scope Analysis",
        0
    )


    # -------------------------------------------------------------
    # Project Summary
    # -------------------------------------------------------------

    summary = data.get(
        "project_summary",
        {}
    )

    doc.add_heading(
        "1. Project Summary",
        level=1
    )

    doc.add_paragraph(
        f"Project Name: "
        f"{summary.get('project_name', 'Unknown')}"
    )

    doc.add_paragraph(
        f"Scope Confidence: "
        f"{summary.get('scope_confidence', 'UNKNOWN')}"
    )

    doc.add_paragraph(
        summary.get(
            "executive_summary",
            ""
        )
    )


    # -------------------------------------------------------------
    # Major Concerns
    # -------------------------------------------------------------

    if summary.get("major_concerns"):

        doc.add_heading(
            "Major Concerns",
            level=2
        )

        for concern in summary["major_concerns"]:

            doc.add_paragraph(
                f"• {concern}"
            )


    # -------------------------------------------------------------
    # Recommended Actions
    # -------------------------------------------------------------

    if summary.get(
        "recommended_next_actions"
    ):

        doc.add_heading(
            "Recommended Next Actions",
            level=2
        )

        for action in summary[
            "recommended_next_actions"
        ]:

            doc.add_paragraph(
                f"• {action}"
            )


    # -------------------------------------------------------------
    # Extracted Scope
    # -------------------------------------------------------------

    doc.add_heading(
        "2. Extracted Scope",
        level=1
    )

    for module in data.get(
        "extracted_scope",
        []
    ):

        doc.add_heading(
            f"Module: "
            f"{module.get('module', 'General')}",
            level=2
        )

        doc.add_paragraph(
            f"Source: "
            f"{module.get('source', 'Unknown')}"
        )

        doc.add_paragraph(
            f"Confidence: "
            f"{module.get('confidence', 'UNKNOWN')}"
        )

        for req in module.get(
            "requirements",
            []
        ):

            doc.add_paragraph(
                f"{req.get('requirement_id', '')} - "
                f"{req.get('description', '')}"
            )

            doc.add_paragraph(
                f"Type: "
                f"{req.get('type', '')}"
            )

            doc.add_paragraph(
                f"Status: "
                f"{req.get('status', '')}"
            )

            doc.add_paragraph(
                f"Source: "
                f"{req.get('source', '')}"
            )


    # -------------------------------------------------------------
    # Dependencies
    # -------------------------------------------------------------

    doc.add_heading(
        "3. Dependencies",
        level=1
    )

    for dependency in data.get(
        "dependencies",
        []
    ):

        doc.add_paragraph(
            f"[{dependency.get('impact', '')}] "
            f"{dependency.get('dependency', '')}"
        )

        doc.add_paragraph(
            dependency.get(
                "description",
                ""
            )
        )

        doc.add_paragraph(
            f"Source: "
            f"{dependency.get('source', '')}"
        )


    # -------------------------------------------------------------
    # Assumptions
    # -------------------------------------------------------------

    doc.add_heading(
        "4. Assumptions",
        level=1
    )

    for assumption in data.get(
        "assumptions",
        []
    ):

        doc.add_paragraph(
            f"[{assumption.get('impact', '')}] "
            f"{assumption.get('assumption', '')}"
        )

        doc.add_paragraph(
            f"Reason: "
            f"{assumption.get('reason', '')}"
        )


    # -------------------------------------------------------------
    # Out of Scope
    # -------------------------------------------------------------

    doc.add_heading(
        "5. Out of Scope",
        level=1
    )

    for item in data.get(
        "out_of_scope",
        []
    ):

        doc.add_paragraph(
            f"• {item.get('item', '')}"
        )

        doc.add_paragraph(
            f"Source: "
            f"{item.get('source', '')}"
        )


    # -------------------------------------------------------------
    # Risks
    # -------------------------------------------------------------

    doc.add_heading(
        "6. Gaps & Risk Audit",
        level=1
    )

    for risk in data.get(
        "gaps_and_risks",
        []
    ):

        doc.add_heading(
            f"[{risk.get('severity', 'INFO')}] "
            f"{risk.get('type', 'Risk')}",
            level=2
        )

        doc.add_paragraph(
            risk.get(
                "description",
                ""
            )
        )

        doc.add_paragraph(
            f"Impact: "
            f"{risk.get('impact', '')}"
        )

        doc.add_paragraph(
            f"Source: "
            f"{risk.get('source', '')}"
        )

        doc.add_paragraph(
            f"Recommended Action: "
            f"{risk.get('recommended_action', '')}"
        )


    # -------------------------------------------------------------
    # Client Questions
    # -------------------------------------------------------------

    doc.add_heading(
        "7. Client Clarification Questions",
        level=1
    )

    for question in data.get(
        "client_questions",
        []
    ):

        doc.add_paragraph(
            f"[{question.get('priority', '')}] "
            f"{question.get('question', '')}"
        )

        doc.add_paragraph(
            f"Reason: "
            f"{question.get('reason', '')}"
        )


    # -------------------------------------------------------------
    # Jira Stories
    # -------------------------------------------------------------

    doc.add_heading(
        "8. Jira User Stories",
        level=1
    )

    for story in data.get(
        "jira_user_stories",
        []
    ):

        doc.add_heading(
            f"{story.get('story_id', '')} - "
            f"{story.get('title', 'User Story')}",
            level=2
        )

        doc.add_paragraph(
            f"As a "
            f"{story.get('user_role', 'User')}, "
            f"I want to "
            f"{story.get('want_statement', '')}, "
            f"so that "
            f"{story.get('so_that_statement', '')}."
        )

        doc.add_paragraph(
            "Source: "
            f"{story.get('source', '')}"
        )

        doc.add_paragraph(
            "Acceptance Criteria:"
        )

        for criterion in story.get(
            "acceptance_criteria",
            []
        ):

            doc.add_paragraph(
                f"• {criterion}"
            )


    # -------------------------------------------------------------
    # Save
    # -------------------------------------------------------------

    target_stream = io.BytesIO()

    doc.save(
        target_stream
    )

    return target_stream.getvalue()


# =============================================================================
# 15. MAIN HEADER BANNER
# =============================================================================

st.markdown(
    """
    <div style="
        background:
            linear-gradient(
                90deg,
                #ec4899 0%,
                #8b5cf6 33%,
                #06b6d4 66%,
                #10b981 100%
            );

        color:#ffffff;

        padding:10px 0px;

        border-radius:10px;

        font-weight:800;

        font-size:1.05rem;

        box-shadow:
            0 0 20px rgba(139,92,246,0.4);

        margin-bottom:24px;
    ">

        <marquee
            behavior="scroll"
            direction="left"
            scrollamount="8"
        >

            ⚡ Smart Scope Handover Engine
            &nbsp;|&nbsp;
            Project Intelligence Layer

            &nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;

            ⚡ Smart Scope Handover Engine
            &nbsp;|&nbsp;
            Project Intelligence Layer

        </marquee>

    </div>
    """,
    unsafe_allow_html=True,
)


# =============================================================================
# 16. TITLE
# =============================================================================

col_title, col_toggle = st.columns(
    [3, 1]
)


with col_title:

    components.html(
        """
        <div style="
            font-family:
                system-ui,
                -apple-system,
                sans-serif;
        ">

            <div style="
                color:#ffffff;
                font-size:2.8rem;
                font-weight:900;
                margin:0;
                line-height:1.2;

                text-shadow:
                    0 0 15px
                    rgba(56,189,248,0.6);
            ">
                Pitch to Project
            </div>

            <div style="
                color:#38bdf8;
                font-size:1.05rem;
                font-weight:700;
                margin-top:6px;
                line-height:1.3;
            ">
                🚀 AI-Powered Scope Intelligence
                & Handover Engine
            </div>

        </div>
        """,
        height=120
    )


with col_toggle:

    demo_mode = st.toggle(
        "Demo Mode (Safe Pitch)",
        value=True
    )


# =============================================================================
# 17. DEMO DATA
# =============================================================================

MOCK_ANALYSIS = {

    "project_summary": {

        "project_name":
            "Smart Scope Handover Engine",

        "executive_summary":
            "The project focuses on automated intake analysis, "
            "scope extraction, risk identification, and Jira "
            "user-story generation.",

        "scope_confidence":
            "MEDIUM",

        "major_concerns": [
            "Operational SLA is not clearly defined.",
            "Document processing limits require confirmation."
        ],

        "recommended_next_actions": [
            "Confirm processing SLA.",
            "Confirm maximum upload size.",
            "Confirm supported user roles."
        ]
    },


    "extracted_scope": [

        {

            "module":
                "AUTHENTICATION & SECURITY",

            "source":
                "Proposal Page 4, Section 2",

            "confidence":
                "HIGH",

            "requirements": [

                {

                    "requirement_id":
                        "REQ-001",

                    "type":
                        "FUNCTIONAL",

                    "description":
                        "Multi-factor authentication "
                        "via SMS or email.",

                    "source":
                        "Proposal Page 4, Section 2",

                    "status":
                        "CONFIRMED"
                },

                {

                    "requirement_id":
                        "REQ-002",

                    "type":
                        "FUNCTIONAL",

                    "description":
                        "Role-based access control "
                        "for Admin and Client roles.",

                    "source":
                        "Proposal Page 4, Section 2",

                    "status":
                        "CONFIRMED"
                }
            ]
        },


        {

            "module":
                "HANDOVER INTELLIGENCE ENGINE",

            "source":
                "Kickoff Transcript Lines 12-45",

            "confidence":
                "HIGH",

            "requirements": [

                {

                    "requirement_id":
                        "REQ-003",

                    "type":
                        "FUNCTIONAL",

                    "description":
                        "Process multiple intake "
                        "document formats.",

                    "source":
                        "Kickoff Transcript Lines 12-45",

                    "status":
                        "CONFIRMED"
                },

                {

                    "requirement_id":
                        "REQ-004",

                    "type":
                        "FUNCTIONAL",

                    "description":
                        "Identify scope gaps and "
                        "generate structured analysis.",

                    "source":
                        "Kickoff Transcript Lines 12-45",

                    "status":
                        "CONFIRMED"
                }
            ]
        }
    ],


    "dependencies": [],

    "assumptions": [],

    "out_of_scope": [],


    "gaps_and_risks": [

        {

            "severity":
                "HIGH",

            "type":
                "Missing SLA",

            "description":
                "No operational SLA is defined "
                "for document processing speed.",

            "impact":
                "Performance expectations cannot "
                "be objectively validated.",

            "source":
                "No SLA found in intake material.",

            "recommended_action":
                "Agree processing-time targets "
                "with the client."
        },


        {

            "severity":
                "MEDIUM",

            "type":
                "Cross-Document Conflict",

            "description":
                "One source requests 500MB upload "
                "support while another specifies "
                "a 200MB limit.",

            "impact":
                "Architecture and infrastructure "
                "capacity may be incorrectly estimated.",

            "source":
                "Kickoff transcript vs SOW.",

            "recommended_action":
                "Confirm the contractual upload limit."
        }
    ],


    "client_questions": [

        {

            "priority":
                "HIGH",

            "question":
                "What is the maximum acceptable "
                "processing time for an uploaded document?",

            "reason":
                "An operational SLA is not defined."
        }
    ],


    "jira_user_stories": [

        {

            "story_id":
                "US-001",

            "title":
                "Configurable Document Processing Engine",

            "user_role":
                "Delivery Lead",

            "want_statement":
                "ingest intake SOWs and transcripts concurrently",

            "so_that_statement":
                "I can automatically extract system "
                "requirements without manual document review",

            "source":
                "Project intake materials",

            "acceptance_criteria": [

                "Given a user uploads supported intake documents",

                "When the user selects Generate Smart Scope",

                "Then the system analyzes the available "
                "requirements and produces structured scope modules."
            ]
        }
    ]
}


# =============================================================================
# 18. MAIN TWO-COLUMN LAYOUT
# =============================================================================

left_col, right_col = st.columns(
    [1, 1],
    gap="medium"
)


# =============================================================================
# 19. LEFT COLUMN
# =============================================================================

with left_col:

    st.header(
        "1. Intake Documents & Media"
    )


    sow_files = st.file_uploader(
        "Proposals / SOWs (.docx)",
        type=["docx"],
        accept_multiple_files=True
    )


    notes_files = st.file_uploader(
        "Transcripts / Notes (.txt)",
        type=["txt"],
        accept_multiple_files=True
    )


    media_files = st.file_uploader(
        "Diagrams & Media (.png, .jpg, .mp4, .mp3)",
        type=[
            "png",
            "jpg",
            "jpeg",
            "mp4",
            "mp3"
        ],
        accept_multiple_files=True
    )


    loose_notes = st.text_area(
        "Or Paste Loose Client Emails / Notes:",

        height=120,

        placeholder=(
            "Paste kickoff notes or client "
            "requirements here..."
        )
    )


    generate_btn = st.button(
        "⚡ GENERATE SMART SCOPE"
    )


# =============================================================================
# 20. RIGHT COLUMN
# =============================================================================

with right_col:

    st.header(
        "2. AI Scope & Handover Analysis"
    )


    # -------------------------------------------------------------
    # GENERATE
    # -------------------------------------------------------------

    if generate_btn:

        progress_card = st.empty()


        # =========================================================
        # DEMO MODE
        # =========================================================

        if demo_mode:

            with progress_card.container(
                border=True
            ):

                st.markdown(
                    "### 🧠 Running Scope Analysis"
                )

                bar_ph = st.empty()


                for i in range(101):

                    time.sleep(
                        0.01
                    )

                    bar_ph.markdown(
                        render_stylish_progress(
                            i,
                            "⚙️ Processing intake materials "
                            "& generating user stories..."
                        ),
                        unsafe_allow_html=True
                    )


                st.success(
                    "✅ Demo Analysis Loaded Successfully!"
                )

                time.sleep(
                    0.4
                )


            progress_card.empty()


            st.session_state[
                "analysis_data"
            ] = MOCK_ANALYSIS


            st.toast(
                "⚡ Demo Analysis Loaded!",
                icon="✅"
            )


        # =========================================================
        # LIVE GEMINI MODE
        # =========================================================

        else:

            raw_text = extract_text_from_uploads(
                sow_files,
                notes_files,
                loose_notes
            )


            if (
                not raw_text.strip()
                and not media_files
            ):

                st.warning(
                    "Please upload at least one "
                    "document/media file or paste "
                    "text before analyzing."
                )

            elif not client:

                st.error(
                    "❌ GEMINI_API_KEY is missing "
                    "from Streamlit Secrets."
                )

            else:

                with progress_card.container(
                    border=True
                ):

                    st.markdown(
                        "### 🧠 Live Gemini Scope Extraction"
                    )

                    bar_ph = st.empty()


                    # -------------------------------------------------
                    # STEP 1
                    # -------------------------------------------------

                    bar_ph.markdown(
                        render_stylish_progress(
                            15,
                            "📄 Step 1/5: Reading "
                            "intake documents..."
                        ),
                        unsafe_allow_html=True
                    )

                    time.sleep(
                        0.3
                    )


                    # -------------------------------------------------
                    # STEP 2
                    # -------------------------------------------------

                    bar_ph.markdown(
                        render_stylish_progress(
                            30,
                            "📊 Step 2/5: Extracting "
                            "paragraphs & tables..."
                        ),
                        unsafe_allow_html=True
                    )

                    time.sleep(
                        0.3
                    )


                    # -------------------------------------------------
                    # STEP 3 - MEDIA
                    # -------------------------------------------------

                    uploaded_media = []


                    if media_files:

                        try:

                            bar_ph.markdown(
                                render_stylish_progress(
                                    45,
                                    "🎞️ Step 3/5: Uploading "
                                    "diagrams & media to Gemini..."
                                ),
                                unsafe_allow_html=True
                            )


                            uploaded_media = (
                                upload_media_to_gemini(
                                    media_files
                                )
                            )


                            uploaded_media = (
                                wait_for_gemini_files(
                                    uploaded_media
                                )
                            )


                        except Exception as exc:

                            st.error(
                                f"❌ Media processing error: "
                                f"{str(exc)}"
                            )

                            uploaded_media = []


                    else:

                        bar_ph.markdown(
                            render_stylish_progress(
                                45,
                                "🎞️ Step 3/5: "
                                "No media files supplied..."
                            ),
                            unsafe_allow_html=True
                        )


                    # -------------------------------------------------
                    # STEP 4
                    # -------------------------------------------------

                    bar_ph.markdown(
                        render_stylish_progress(
                            65,
                            "⚡ Step 4/5: Sending "
                            "project evidence to Gemini..."
                        ),
                        unsafe_allow_html=True
                    )

                    time.sleep(
                        0.3
                    )


                    # -------------------------------------------------
                    # STEP 5
                    # -------------------------------------------------

                    bar_ph.markdown(
                        render_stylish_progress(
                            80,
                            "🔍 Step 5/5: Auditing scope, "
                            "risks, gaps & Jira stories..."
                        ),
                        unsafe_allow_html=True
                    )


                    # -------------------------------------------------
                    # AI CALL
                    # -------------------------------------------------

                    result = analyze_with_gemini(
                        raw_text,
                        uploaded_media
                    )


                    if result:

                        bar_ph.markdown(
                            render_stylish_progress(
                                100,
                                "✅ Smart Scope Handover "
                                "Analysis Complete!"
                            ),
                            unsafe_allow_html=True
                        )


                        st.success(
                            "✅ Scope Successfully Extracted!"
                        )


                        time.sleep(
                            0.5
                        )


                        st.session_state[
                            "analysis_data"
                        ] = result


                        st.toast(
                            "⚡ Live Gemini Extraction Complete!",
                            icon="✅"
                        )


                    else:

                        st.error(
                            "❌ Extraction Failed"
                        )


                progress_card.empty()


    # =============================================================================
    # 21. DEFAULT DATA
    # =============================================================================

    if "analysis_data" not in st.session_state:

        st.session_state[
            "analysis_data"
        ] = MOCK_ANALYSIS


    data = st.session_state[
        "analysis_data"
    ]


    # =============================================================================
    # 22. TABS
    # =============================================================================

    tab_scope, tab_risks, tab_jira = st.tabs(
        [
            "📌 Extracted Scope",
            "🚨 Missing Items & Risks",
            "🚀 Jira User Stories"
        ]
    )


    # =============================================================================
    # 23. SCOPE TAB
    # =============================================================================

    with tab_scope:

        for item in data.get(
            "extracted_scope",
            []
        ):

            with st.container(
                border=True
            ):

                st.markdown(
                    f"### 📌 MODULE: "
                    f"{item.get('module', 'General Scope')}"
                )


                st.caption(
                    f"Source: "
                    f"{item.get('source', 'Uploaded Files')}"
                )


                st.caption(
                    f"Confidence: "
                    f"{item.get('confidence', 'UNKNOWN')}"
                )


                # -----------------------------------------------------
                # New format
                # -----------------------------------------------------

                requirements = item.get(
                    "requirements",
                    []
                )


                if requirements:

                    for req in requirements:

                        st.markdown(
                            f"**{req.get('requirement_id', '')}** "
                            f"• {req.get('description', '')}"
                        )

                        st.markdown(
                            f"`{req.get('type', '')}` "
                            f"`{req.get('status', '')}`"
                        )

                        if req.get("source"):

                            st.caption(
                                f"Source: "
                                f"{req.get('source')}"
                            )


                # -----------------------------------------------------
                # Backward compatibility with old mock structure
                # -----------------------------------------------------

                for pt in item.get(
                    "points",
                    []
                ):

                    st.markdown(
                        f"• {pt}"
                    )


    # =============================================================================
    # 24. RISKS TAB
    # =============================================================================

    with tab_risks:

        for risk in data.get(
            "gaps_and_risks",
            []
        ):

            with st.container(
                border=True
            ):

                severity = (
                    risk.get(
                        "severity",
                        "INFO"
                    )
                    .upper()
                )


                badge_color = (
                    "🔴"
                    if severity == "HIGH"
                    else
                    "🟡"
                    if severity == "MEDIUM"
                    else
                    "🔵"
                )


                st.markdown(
                    f"### {badge_color} "
                    f"[{severity}] "
                    f"{risk.get('type', 'Risk')}"
                )


                st.write(
                    risk.get(
                        "description",
                        ""
                    )
                )


                if risk.get("impact"):

                    st.markdown(
                        f"**Impact:** "
                        f"{risk.get('impact')}"
                    )


                if risk.get("source"):

                    st.caption(
                        f"Source: "
                        f"{risk.get('source')}"
                    )


                if risk.get(
                    "recommended_action"
                ):

                    st.markdown(
                        f"**Recommended Action:** "
                        f"{risk.get('recommended_action')}"
                    )


        # -------------------------------------------------------------
        # Client Questions
        # -------------------------------------------------------------

        questions = data.get(
            "client_questions",
            []
        )


        if questions:

            st.markdown(
                "### ❓ Client Clarification Questions"
            )


            for question in questions:

                with st.container(
                    border=True
                ):

                    st.markdown(
                        f"**[{question.get('priority', 'MEDIUM')}]** "
                        f"{question.get('question', '')}"
                    )

                    st.caption(
                        f"Reason: "
                        f"{question.get('reason', '')}"
                    )


    # =============================================================================
    # 25. JIRA TAB
    # =============================================================================

    with tab_jira:

        for story in data.get(
            "jira_user_stories",
            []
        ):

            with st.container(
                border=True
            ):

                st.markdown(
                    f"### 🚀 "
                    f"{story.get('story_id', '')} "
                    f"{story.get('title', 'User Story')}"
                )


                st.markdown(
                    f"**As a** "
                    f"`{story.get('user_role', 'User')}`, "
                    f"**I want to** "
                    f"{story.get('want_statement', '')} "
                    f"**so that** "
                    f"{story.get('so_that_statement', '')}."
                )


                if story.get("source"):

                    st.caption(
                        f"Source: "
                        f"{story.get('source')}"
                    )


                st.markdown(
                    "**Acceptance Criteria:**"
                )


                for ac in story.get(
                    "acceptance_criteria",
                    []
                ):

                    st.markdown(
                        f"- `{ac}`"
                    )


    # =============================================================================
    # 26. REPORT DOWNLOAD
    # =============================================================================

    st.divider()


    docx_bytes = build_docx_report(
        data
    )


    st.download_button(

        label=
            "📄 Export Handover Report (.docx)",

        data=
            docx_bytes,

        file_name=
            "Pitch_to_Project_Handover_Report.docx",

        mime=
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document",

        use_container_width=True
    )
